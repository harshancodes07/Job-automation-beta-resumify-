from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


def generate_docx(tailored_text: str) -> BytesIO:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    lines = tailored_text.strip().split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(stripped[2:])
            p.style = doc.styles["Heading 1"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:])
            p.style = doc.styles["Heading 2"]
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_pdf(tailored_text: str) -> BytesIO:
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=54, bottomMargin=54)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ResumeHeading",
        parent=styles["Heading1"],
        fontSize=16,
        spaceAfter=8,
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="SectionHeading",
        parent=styles["Heading2"],
        fontSize=12,
        spaceAfter=6,
        spaceBefore=12,
    ))
    styles.add(ParagraphStyle(
        name="ResumeBody",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=4,
        leading=14,
    ))
    styles.add(ParagraphStyle(
        name="BulletItem",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=2,
        leading=14,
        leftIndent=20,
        bulletIndent=10,
    ))

    elements = []
    lines = tailored_text.strip().split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            elements.append(Spacer(1, 6))
            continue

        escaped = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if stripped.startswith("# "):
            elements.append(Paragraph(escaped[2:], styles["ResumeHeading"]))
        elif stripped.startswith("## "):
            elements.append(Paragraph(escaped[3:], styles["SectionHeading"]))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            elements.append(Paragraph(f"• {escaped[2:]}", styles["BulletItem"]))
        else:
            elements.append(Paragraph(escaped, styles["ResumeBody"]))

    doc.build(elements)
    buf.seek(0)
    return buf
